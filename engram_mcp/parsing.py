from __future__ import annotations

import fnmatch
import hashlib
import logging
import os
from contextlib import ExitStack
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterable, Iterator, List, Optional, Tuple

from .chunking import Chunk, chunk_lines, chunk_text, make_chunk_id, token_count
from .security import PathContext


SUPPORTED_TEXT_EXTS = {
    ".txt", ".md", ".rst", ".py", ".js", ".ts", ".tsx", ".jsx", ".cs", ".java", ".go", ".rs", ".cpp", ".c", ".h", ".hpp", ".sql", ".json", ".yaml", ".yml", ".toml", ".ini", ".cfg", ".conf", ".html", ".css", ".scss", ".xml", ".sh", ".bash", ".ps1",
}

SUPPORTED_DOC_EXTS = {".pdf", ".docx"}


# ---------------------------------------------------------------------------
# Graph dataclasses
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class GraphNodeObj:
    """A single code-graph node (function / class) extracted from a source file."""

    node_id: str  # Unique within a project:  "path:type:name"
    node_type: str  # "function" | "class"
    name: str  # Simple identifier; may include qualifier for C++ (MyClass::method)
    file_path: str  # Relative path within the project root
    start_line: int  # 1-based, inclusive
    end_line: int  # 1-based, inclusive
    metadata: Dict[str, Any]  # e.g. {"language": "rust"}


@dataclass(frozen=True)
class GraphEdgeObj:
    """A directed edge in the code graph."""

    source_id: str
    target_id: str
    edge_type: str  # "calls", "contains", "imports", …


# ---------------------------------------------------------------------------
# Polyglot query registry
# ---------------------------------------------------------------------------
# Maps file extension → (tree-sitter language name, SCM query string).
# Adding a new language only requires a new entry here.

LANGUAGE_CONFIG: Dict[str, Tuple[str, str]] = {
    # --- Systems languages ---
    ".rs": (
        "rust",
        """
        (function_item name: (identifier) @name) @func
        (struct_item name: (type_identifier) @name) @cls
        (impl_item type: (type_identifier) @name) @impl
        """,
    ),
    ".go": (
        "go",
        """
        (function_declaration name: (identifier) @name) @func
        (method_declaration name: (field_identifier) @name) @method
        (type_declaration (type_spec name: (type_identifier) @name)) @cls
        """,
    ),
    ".c": (
        "c",
        """
        (function_definition) @func
        (struct_specifier name: (type_identifier) @name) @cls
        """,
    ),
    ".cpp": (
        "cpp",
        """
        (function_definition) @func
        (class_specifier name: (type_identifier) @name) @cls
        """,
    ),
    ".h": (
        "cpp",
        """
        (class_specifier name: (type_identifier) @name) @cls
        """,
    ),
    ".hpp": (
        "cpp",
        """
        (class_specifier name: (type_identifier) @name) @cls
        """,
    ),
    # --- Enterprise languages ---
    ".java": (
        "java",
        """
        (method_declaration name: (identifier) @name) @func
        (class_declaration name: (identifier) @name) @cls
        (interface_declaration name: (identifier) @name) @cls
        """,
    ),
    ".cs": (
        "c_sharp",
        """
        (method_declaration name: (identifier) @name) @func
        (class_declaration name: (identifier) @name) @cls
        (interface_declaration name: (identifier) @name) @cls
        """,
    ),
    # --- Web languages ---
    ".js": (
        "javascript",
        """
        (function_declaration name: (identifier) @name) @func
        (class_declaration name: (identifier) @name) @cls
        """,
    ),
    ".ts": (
        "typescript",
        """
        (function_declaration name: (identifier) @name) @func
        (class_declaration name: (type_identifier) @name) @cls
        (interface_declaration name: (type_identifier) @name) @cls
        """,
    ),
    ".tsx": (
        "tsx",
        """
        (function_declaration name: (identifier) @name) @func
        (class_declaration name: (type_identifier) @name) @cls
        """,
    ),
}


# ---------------------------------------------------------------------------
# Graph extraction helpers
# ---------------------------------------------------------------------------


def _extract_python_graph(
    content: str, file_path: str
) -> Tuple[List[GraphNodeObj], List[GraphEdgeObj]]:
    """Extract functions and classes from Python source using the stdlib ast."""
    import ast as _ast

    try:
        tree = _ast.parse(content)
    except SyntaxError:
        return [], []

    nodes: List[GraphNodeObj] = []
    for node in _ast.walk(tree):
        if isinstance(node, (_ast.FunctionDef, _ast.AsyncFunctionDef)):
            nodes.append(
                GraphNodeObj(
                    node_id=f"{file_path}:function:{node.name}",
                    node_type="function",
                    name=node.name,
                    file_path=file_path,
                    start_line=node.lineno,
                    end_line=node.end_lineno or node.lineno,
                    metadata={"language": "python"},
                )
            )
        elif isinstance(node, _ast.ClassDef):
            nodes.append(
                GraphNodeObj(
                    node_id=f"{file_path}:class:{node.name}",
                    node_type="class",
                    name=node.name,
                    file_path=file_path,
                    start_line=node.lineno,
                    end_line=node.end_lineno or node.lineno,
                    metadata={"language": "python"},
                )
            )
    return nodes, []


def _get_node_name(node: Any, content: str) -> "Optional[str]":
    """Navigate tree-sitter grammar quirks to extract the definition name.

    Most grammars expose a direct ``name`` field on the definition node.
    Two exceptions need special handling:

    * C / C++ ``function_definition`` – the identifier is nested inside a
      (possibly pointer-wrapped) ``function_declarator``.
    * Go ``type_declaration`` – the ``name`` field lives on the inner
      ``type_spec`` child, not on the declaration node itself.
    """
    # Fast path – works for Rust, Go func/method, Java, C#, JS, TS, C struct,
    # C++ class, …
    name_node = node.child_by_field_name("name")
    if name_node:
        return content[name_node.start_byte : name_node.end_byte]

    # C / C++ function_definition: unwrap the declarator chain.
    # function_definition → (pointer_declarator →)* function_declarator → identifier
    if node.type == "function_definition":
        decl = node.child_by_field_name("declarator")
        while decl:
            if decl.type in (
                "function_declarator",
                "pointer_declarator",
                "reference_declarator",
            ):
                decl = decl.child_by_field_name("declarator")
            else:
                # Terminal: identifier, field_identifier, or scope_resolution
                return content[decl.start_byte : decl.end_byte]

    # Go type_declaration: the name is on the inner type_spec node.
    if node.type == "type_declaration":
        for child in node.children:
            if child.type == "type_spec":
                spec_name = child.child_by_field_name("name")
                if spec_name:
                    return content[spec_name.start_byte : spec_name.end_byte]

    return None


def _extract_polyglot_graph(
    content: str, file_path: str, ext: str
) -> Tuple[List[GraphNodeObj], List[GraphEdgeObj]]:
    """Extract functions and classes from any language in LANGUAGE_CONFIG.

    Uses tree-sitter-languages for parsing.  Returns empty lists when the
    dependency is missing or the language is unavailable — never raises.
    """
    config = LANGUAGE_CONFIG.get(ext)
    if not config:
        return [], []

    lang_name, query_scm = config

    try:
        from tree_sitter_languages import get_language, get_parser

        language = get_language(lang_name)
        parser = get_parser(lang_name)
    except ImportError:
        logging.warning("tree-sitter-languages not installed; graph extraction skipped.")
        return [], []
    except Exception:
        logging.warning(
            "Language %s not available in tree-sitter-languages.", lang_name, exc_info=True
        )
        return [], []

    tree = parser.parse(content.encode("utf-8"))
    try:
        query = language.query(query_scm)
    except Exception:
        logging.warning("Invalid tree-sitter query for %s.", lang_name, exc_info=True)
        return [], []

    raw_captures = query.captures(tree.root_node)
    # Normalise across tree-sitter versions:
    #   < 0.21  →  [(Node, str), …]
    #   >= 0.21 →  {str: [Node, …]}
    if isinstance(raw_captures, dict):
        captures: List[Tuple[Any, str]] = [
            (n, name) for name, nodes in raw_captures.items() for n in nodes
        ]
    else:
        captures = list(raw_captures)

    nodes: List[GraphNodeObj] = []
    for node, tag in captures:
        # @name captures are only used as markers in the SCM query; the actual
        # name is extracted via _get_node_name on the parent definition node.
        if tag == "name":
            continue

        node_type = "function"
        if tag in ("cls", "interface", "struct"):
            node_type = "class"

        name = _get_node_name(node, content)
        if not name:
            continue  # anonymous / un-navigable (e.g. Rust impl blocks)

        nodes.append(
            GraphNodeObj(
                node_id=f"{file_path}:{node_type}:{name}",
                node_type=node_type,
                name=name,
                file_path=file_path,
                start_line=node.start_point[0] + 1,
                end_line=node.end_point[0] + 1,
                metadata={"language": lang_name},
            )
        )
    return nodes, []


def _is_ignored(path: str, ignore_patterns: List[str]) -> bool:
    basename = os.path.basename(path)
    for pat in ignore_patterns:
        if fnmatch.fnmatch(path, pat) or fnmatch.fnmatch(basename, pat):
            return True
        # fnmatch does not honour the "any directory depth" semantics of **;
        # strip leading **/ prefixes so that patterns like **/*.log also
        # match files directly at the root (where the relative path has no
        # directory component).
        stripped = pat
        while stripped.startswith("**/"):
            stripped = stripped[3:]
        if stripped != pat and fnmatch.fnmatch(basename, stripped):
            return True
    return False


def iter_files(path_context: PathContext, root: str, ignore_patterns: List[str]) -> Iterator[Path]:
    root_path = path_context.resolve_path(root)
    for p in path_context.iter_files(root_path):
        if not p.is_file():
            continue
        rel = str(p.relative_to(root_path))
        if _is_ignored(rel, ignore_patterns):
            continue
        yield p


def read_text_streaming(
    path_context: PathContext,
    path: Path,
    root: Path,
    max_bytes: int,
    *,
    chunk_size: int = 1024 * 1024,
) -> Iterable[str]:
    # Avoid reading enormous files at once.
    size = path_context.stat(path).st_size
    if size > max_bytes:
        raise ValueError(f"File too large: {path} ({size} bytes)")
    with ExitStack() as stack:
        handle = stack.enter_context(
            path_context.open_file(path, "r", encoding="utf-8", errors="ignore")
        )
        while True:
            chunk = handle.read(chunk_size)
            if not chunk:
                break
            yield chunk


def hash_file(path_context: PathContext, path: Path, root: Path, max_bytes: int) -> str:
    resolved = path_context.resolve_path(path)
    flags = os.O_RDONLY
    if hasattr(os, "O_NOFOLLOW"):
        flags |= os.O_NOFOLLOW
    if hasattr(os, "O_CLOEXEC"):
        flags |= os.O_CLOEXEC
    fd = os.open(resolved, flags)
    try:
        st_before = os.fstat(fd)
        if st_before.st_size > max_bytes:
            raise ValueError(f"File too large: {path} ({st_before.st_size} bytes)")
        h = hashlib.sha256()
        with os.fdopen(fd, "rb", closefd=False) as handle:
            while True:
                chunk = handle.read(1024 * 1024)
                if not chunk:
                    break
                h.update(chunk)
        st_after = os.fstat(fd)
        if st_after.st_mtime_ns != st_before.st_mtime_ns or st_after.st_size != st_before.st_size:
            raise RuntimeError(f"File changed during hashing: {path}")
        return h.hexdigest()
    finally:
        os.close(fd)


def parse_file_to_chunks(
    *,
    path_context: PathContext,
    path: Path,
    root: str,
    project_type: str,
    chunk_size_tokens: int,
    overlap_tokens: int,
    max_file_size_mb: int,
) -> List[Chunk]:
    root_path = path_context.resolve_path(root)
    if root_path not in path.parents and path != root_path:
        raise ValueError(f"Path escapes root: {path}")
    ext = path.suffix.lower()
    rel = str(path.relative_to(root))
    base_meta: Dict[str, object] = {
        "item_name": rel,
        "path": str(path),
        "ext": ext,
        "type": "file",
        "project_type": project_type,
    }

    max_bytes = int(max_file_size_mb) * 1024 * 1024

    if ext in SUPPORTED_TEXT_EXTS:
        base_id = make_chunk_id(str(path))
        return chunk_lines(
            lines=read_text_streaming(path_context, path, root=Path(root), max_bytes=max_bytes),
            base_id=base_id,
            meta=base_meta,
            target_tokens=chunk_size_tokens,
            overlap_tokens=overlap_tokens,
        )

    if ext == ".pdf":
        # Streaming: page by page
        from pypdf import PdfReader

        # Check PDF file size before loading (prevent OOM/DoS)
        size = path_context.stat(path).st_size
        if size > max_bytes:
            raise ValueError(f"PDF file too large: {path} ({size} bytes)")

        max_text_chars = max_bytes * 4
        chunks: List[Chunk] = []
        base_id = make_chunk_id(str(path))
        total_chars = 0
        with ExitStack() as stack:
            handle = stack.enter_context(path_context.open_file(path, "rb"))
            reader = PdfReader(handle)
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text() or ""
                page_len = len(page_text)
                if total_chars + page_len > max_text_chars:
                    raise ValueError(f"PDF extracted text too large: {path}")
                total_chars += page_len
                if not page_text.strip():
                    continue
                page_meta = {**base_meta, "type": "pdf_page", "page": i + 1}
                # Keep pages as their own units, then chunk if needed
                if token_count(page_text) <= chunk_size_tokens * 2:
                    cid = make_chunk_id(base_id, f"page:{i+1}", page_text)
                    chunks.append(Chunk(cid, page_text, token_count(page_text), page_meta))
                else:
                    chunks.extend(
                        chunk_text(
                            text=page_text,
                            base_id=make_chunk_id(base_id, f"page:{i+1}"),
                            meta=page_meta,
                            target_tokens=chunk_size_tokens,
                            overlap_tokens=overlap_tokens,
                        )
                    )
        return chunks

    if ext == ".docx":
        from docx import Document

        max_text_chars = max_bytes * 4
        paras: List[str] = []
        total_chars = 0
        with ExitStack() as stack:
            handle = stack.enter_context(path_context.open_file(path, "rb"))
            doc = Document(handle)
            for p in doc.paragraphs:
                if not p.text or not p.text.strip():
                    continue
                total_chars += len(p.text)
                if total_chars > max_text_chars:
                    raise ValueError(f"DOCX extracted text too large: {path}")
                paras.append(p.text)
        text = "\n".join(paras)
        base_id = make_chunk_id(str(path))
        return chunk_text(
            text=text,
            base_id=base_id,
            meta={**base_meta, "type": "docx"},
            target_tokens=chunk_size_tokens,
            overlap_tokens=overlap_tokens,
        )

    return []


def hash_and_parse_file(
    *,
    path_context: PathContext,
    path: Path,
    root: str,
    project_type: str,
    chunk_size_tokens: int,
    overlap_tokens: int,
    max_file_size_mb: int,
) -> "tuple[List[Chunk], str, List[GraphNodeObj], List[GraphEdgeObj]]":
    """Stream a file to return its chunks, SHA-256 hash, and graph nodes/edges.

    Text files are parsed line-by-line while hashing the raw bytes to avoid
    large in-memory buffers. Binary formats (PDF/DOCX) fall back to a second
    read via their parsers after hashing, prioritizing memory safety over
    single-read I/O.
    """
    import codecs

    max_bytes = int(max_file_size_mb) * 1024 * 1024
    root_path = path_context.resolve_path(root)
    if root_path not in path.parents and path != root_path:
        raise ValueError(f"Path escapes root: {path}")

    size = path_context.stat(path).st_size
    if size > max_bytes:
        raise ValueError(f"File too large: {path} ({size} bytes)")

    ext = path.suffix.lower()
    rel = str(path.relative_to(root))
    base_meta: Dict[str, object] = {
        "item_name": rel,
        "path": str(path),
        "ext": ext,
        "type": "file",
        "project_type": project_type,
    }

    if ext in SUPPORTED_TEXT_EXTS:
        hasher = hashlib.sha256()
        decoder = codecs.getincrementaldecoder("utf-8")(errors="ignore")
        buffer = ""
        # Accumulate decoded text only for extensions that need graph extraction
        # so that large plain-text files (README, SQL, …) stay streaming.
        _needs_graph = (ext == ".py") or (ext in LANGUAGE_CONFIG)
        _content_parts: List[str] = []

        def _iter_lines() -> Iterator[str]:
            nonlocal buffer
            with path_context.open_file(path, "rb") as fh:
                while True:
                    chunk = fh.read(1024 * 1024)
                    if not chunk:
                        break
                    hasher.update(chunk)
                    text_chunk = decoder.decode(chunk)
                    if not text_chunk:
                        continue
                    if _needs_graph:
                        _content_parts.append(text_chunk)
                    buffer += text_chunk
                    lines = buffer.splitlines(keepends=True)
                    if lines:
                        if not buffer.endswith(("\n", "\r")):
                            buffer = lines.pop() if lines else buffer
                        else:
                            buffer = ""
                        for line in lines:
                            yield line
                # Flush any incomplete multi-byte sequence at EOF.
                final_flush = decoder.decode(b"", final=True)
                if _needs_graph and final_flush:
                    _content_parts.append(final_flush)
                tail = buffer + final_flush
                if tail:
                    for line in tail.splitlines(keepends=True):
                        yield line

        base_id = make_chunk_id(str(path))
        chunks = chunk_lines(
            lines=_iter_lines(),
            base_id=base_id,
            meta=base_meta,
            target_tokens=chunk_size_tokens,
            overlap_tokens=overlap_tokens,
        )

        # --- graph extraction (single-pass: content already buffered) ---
        nodes: List[GraphNodeObj] = []
        edges: List[GraphEdgeObj] = []
        if _needs_graph:
            content_str = "".join(_content_parts)
            if ext == ".py":
                nodes, edges = _extract_python_graph(content_str, rel)
            elif ext in LANGUAGE_CONFIG:
                nodes, edges = _extract_polyglot_graph(content_str, rel, ext)
        return chunks, hasher.hexdigest(), nodes, edges

    if ext == ".pdf":
        import io
        from pypdf import PdfReader

        # Single read: hash and parse from the same buffer.  File size is
        # already validated against max_bytes above, so loading into memory
        # is safe and avoids a redundant disk pass.
        with path_context.open_file(path, "rb") as fh:
            raw = fh.read()
        content_hash = hashlib.sha256(raw).hexdigest()
        max_text_chars = max_bytes * 4
        reader = PdfReader(io.BytesIO(raw))
        del raw  # release the raw buffer; pages are decoded on demand
        chunks: List[Chunk] = []
        base_id = make_chunk_id(str(path))
        total_chars = 0
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            page_len = len(page_text)
            if total_chars + page_len > max_text_chars:
                raise ValueError(f"PDF extracted text too large: {path}")
            total_chars += page_len
            if not page_text.strip():
                continue
            page_meta = {**base_meta, "type": "pdf_page", "page": i + 1}
            if token_count(page_text) <= chunk_size_tokens * 2:
                cid = make_chunk_id(base_id, f"page:{i+1}", page_text)
                chunks.append(Chunk(cid, page_text, token_count(page_text), page_meta))
            else:
                chunks.extend(
                    chunk_text(
                        text=page_text,
                        base_id=make_chunk_id(base_id, f"page:{i+1}"),
                        meta=page_meta,
                        target_tokens=chunk_size_tokens,
                        overlap_tokens=overlap_tokens,
                    )
                )
        return chunks, content_hash, [], []

    if ext == ".docx":
        import io
        from docx import Document

        # Single read: hash and parse from the same buffer.
        with path_context.open_file(path, "rb") as fh:
            raw = fh.read()
        content_hash = hashlib.sha256(raw).hexdigest()
        max_text_chars = max_bytes * 4
        doc = Document(io.BytesIO(raw))
        del raw
        paras: List[str] = []
        total_chars = 0
        for p in doc.paragraphs:
            if not p.text or not p.text.strip():
                continue
            total_chars += len(p.text)
            if total_chars > max_text_chars:
                raise ValueError(f"DOCX extracted text too large: {path}")
            paras.append(p.text)
        text = "\n".join(paras)
        base_id = make_chunk_id(str(path))
        chunks = chunk_text(
            text=text,
            base_id=base_id,
            meta={**base_meta, "type": "docx"},
            target_tokens=chunk_size_tokens,
            overlap_tokens=overlap_tokens,
        )
        return chunks, content_hash, [], []

    content_hash = hash_file(path_context, path, Path(root), max_bytes)
    return [], content_hash, [], []


@dataclass(frozen=True)
class ParsedFile:
    file_path: str
    mtime_ns: int
    size_bytes: int
    content_hash: str
    chunks: List[Chunk]
    nodes: List[GraphNodeObj] = field(default_factory=list)
    edges: List[GraphEdgeObj] = field(default_factory=list)
